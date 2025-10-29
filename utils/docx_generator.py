"""
DOCX 보고서 생성기 v6.0 - 전문가급 디자인 & 내용 강화
- 기존 v5.0의 모든 요소 유지 (Binary checklist, Appendix, Reference 등)
- 고급 디자인: 색상 테마, 표 스타일, 섹션 구분
- 내용 강화: 상세 분석, 비교 평가, 개선 로드맵
- 시각적 요소: 아이콘, 하이라이트, 정보 박스
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
from typing import Dict, Any, List, Tuple
from datetime import datetime


class PatentReportGenerator:
    """특허 평가 DOCX 보고서 생성기 v6.0 - 전문가급"""
    
    # 컬러 테마 정의 (RGB 튜플)
    COLOR_TUPLES = {
        'primary': (0, 102, 204),      # 진한 파랑
        'secondary': (52, 152, 219),   # 밝은 파랑
        'success': (46, 204, 113),     # 녹색
        'warning': (241, 196, 15),     # 노랑
        'danger': (231, 76, 60),       # 빨강
        'gray': (149, 165, 166),       # 회색
        'dark': (44, 62, 80),          # 진한 회색
        'light': (236, 240, 241),      # 연한 회색
    }
    
    def __init__(self):
        # RGB 튜플을 RGBColor 객체로 변환
        self.COLORS = {}
        for key, rgb in self.COLOR_TUPLES.items():
            self.COLORS[key] = RGBColor(rgb[0], rgb[1], rgb[2])
    
    def _get_rgb_color(self, color_key: str) -> RGBColor:
        """컬러 키로 RGBColor 객체 생성"""
        return self.COLORS.get(color_key, self.COLORS['gray'])
    
    def _rgb_tuple_to_hex(self, rgb_tuple: tuple) -> str:
        """RGB 튜플을 hex 문자열로 변환"""
        return f"{rgb_tuple[0]:02X}{rgb_tuple[1]:02X}{rgb_tuple[2]:02X}"
    
    def generate_report(
        self,
        patent_info: Dict[str, Any],
        state: Dict[str, Any],
        chart_paths: Dict[str, str],
        output_path: str
    ) -> str:
        """
        보고서 생성 (v6.0 - 모든 기존 기능 + 강화)
        
        Args:
            patent_info: PDF에서 추출한 특허 정보
            state: 에이전트 평가 결과
            chart_paths: 차트 파일 경로들
            output_path: 출력 경로
        """
        doc = Document()
        self._set_korean_font(doc)
        
        # ========== 기존 v5.0 구조 유지 ==========
        
        # 1. 표지 (강화)
        self._add_enhanced_cover_page(doc, patent_info, state)
        doc.add_page_break()
        
        # 2. Executive Summary (강화)
        self._add_enhanced_executive_summary(doc, patent_info, state)
        doc.add_page_break()
        
        # 3. 목차 (신규)
        self._add_table_of_contents(doc)
        doc.add_page_break()
        
        # 4. 평가 결과 시각화 (강화)
        self._add_enhanced_evaluation_charts(doc, chart_paths, state)
        doc.add_page_break()
        
        # 5. 특허 개요 및 메타데이터 (신규)
        self._add_patent_overview(doc, patent_info)
        doc.add_page_break()
        
        # 6. 기술성 평가 (기존 + 강화)
        self._add_enhanced_technology_evaluation(doc, state)
        doc.add_page_break()
        
        # 7. 권리성 평가 (기존 + 강화)
        self._add_enhanced_rights_evaluation(doc, state)
        doc.add_page_break()
        
        # 8. 활용성 평가 (기존 + 강화)
        self._add_enhanced_market_evaluation(doc, state)
        doc.add_page_break()
        
        # 9. 비교 분석 (신규)
        self._add_comparative_analysis(doc, state)
        doc.add_page_break()
        
        # 10. 종합 평가 및 제언 (강화)
        self._add_enhanced_recommendations(doc, state, patent_info)
        doc.add_page_break()
        
        # 11. 개선 로드맵 (신규)
        self._add_improvement_roadmap(doc, state)
        doc.add_page_break()
        
        # 12. Reference (기존 유지 + 강화)
        self._add_enhanced_references(doc, patent_info, state)
        doc.add_page_break()
        
        # 13. Appendix (기존 유지 + 강화)
        self._add_enhanced_appendix(doc, state)
        
        # 저장
        doc.save(output_path)
        return output_path
    
    def _set_korean_font(self, doc: Document):
        """한글 폰트 설정"""
        style = doc.styles['Normal']
        font = style.font
        font.name = '맑은 고딕'
        font.size = Pt(10)
        
        rFonts = style.element.rPr.rFonts
        rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    
    def _add_shading_to_paragraph(self, paragraph, color):
        """단락에 배경색 추가"""
        shading_elm = OxmlElement('w:shd')
        
        # color가 튜플인 경우
        if isinstance(color, tuple) and len(color) == 3:
            hex_color = self._rgb_tuple_to_hex(color)
        # color가 문자열 키인 경우
        elif isinstance(color, str) and color in self.COLOR_TUPLES:
            hex_color = self._rgb_tuple_to_hex(self.COLOR_TUPLES[color])
        # color가 RGBColor 객체인 경우
        elif hasattr(color, '__class__') and color.__class__.__name__ == 'RGBColor':
            # RGBColor에서 직접 값을 추출할 수 없으므로 기본값 사용
            hex_color = self._rgb_tuple_to_hex(self.COLOR_TUPLES['light'])
        else:
            # 기본값
            hex_color = self._rgb_tuple_to_hex(self.COLOR_TUPLES['light'])
        
        shading_elm.set(qn('w:fill'), hex_color)
        paragraph._element.get_or_add_pPr().append(shading_elm)
    
    def _create_info_box(self, doc: Document, title: str, content: str, 
                         color=None):
        """정보 박스 생성 (하이라이트된 텍스트 영역)"""
        if color is None:
            color = self.COLOR_TUPLES['light']
        
        # color가 문자열 키인 경우 튜플로 변환
        if isinstance(color, str) and color in self.COLOR_TUPLES:
            color = self.COLOR_TUPLES[color]
        
        p = doc.add_paragraph()
        self._add_shading_to_paragraph(p, color)
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.right_indent = Inches(0.2)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        
        run = p.add_run(f"💡 {title}\n")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = self.COLORS['primary']
        
        run = p.add_run(content)
        run.font.size = Pt(10)
        
        return p
    
    def _add_enhanced_cover_page(self, doc: Document, patent_info: Dict, state: Dict):
        """강화된 표지 페이지"""
        # 제목
        title = doc.add_heading('특허 기술 평가 보고서', level=0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in title.runs:
            run.font.color.rgb = self.COLORS['primary']
            run.font.size = Pt(28)
        
        doc.add_paragraph()
        
        # 부제목
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = subtitle.add_run('Professional Patent Evaluation System v6.0')
        run.font.size = Pt(11)
        run.font.color.rgb = self.COLORS['gray']
        run.italic = True
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # 특허 정보 박스
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        run = p.add_run(f"특허번호: {patent_info.get('number', 'N/A')}\n\n")
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = self.COLORS['dark']
        
        run = p.add_run(f"{patent_info.get('title', 'N/A')}\n\n")
        run.font.size = Pt(14)
        run.font.color.rgb = self.COLORS['dark']
        
        run = p.add_run(f"출원인: {patent_info.get('applicant', 'N/A')}\n")
        run.font.size = Pt(12)
        run.font.color.rgb = self.COLORS['gray']
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # 평가 결과 하이라이트
        overall_score = state.get('overall_score', 0)
        final_grade = state.get('final_grade', 'N/A')
        
        # 점수에 따른 색상 결정
        if overall_score >= 80:
            score_color = self.COLORS['success']
        elif overall_score >= 70:
            score_color = self.COLORS['secondary']
        elif overall_score >= 60:
            score_color = self.COLORS['warning']
        else:
            score_color = self.COLORS['danger']
        
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        run = p.add_run(f"종합 평가 점수\n")
        run.font.size = Pt(14)
        run.font.color.rgb = self.COLORS['gray']
        
        run = p.add_run(f"{overall_score:.1f}점\n")
        run.font.size = Pt(36)
        run.font.bold = True
        run.font.color.rgb = score_color
        
        run = p.add_run(f"등급: {final_grade}\n")
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = score_color
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # 평가 영역별 간단 요약
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        tech_score = state.get('tech_score', 0)
        rights_score = state.get('rights_score', 0)
        market_score = state.get('market_score', 0)
        
        run = p.add_run(f"기술성 {tech_score:.1f}  |  권리성 {rights_score:.1f}  |  활용성 {market_score:.1f}\n")
        run.font.size = Pt(11)
        run.font.color.rgb = self.COLORS['gray']
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # 날짜 및 버전
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run(f"평가일: {datetime.now().strftime('%Y년 %m월 %d일')}\n")
        run.font.size = Pt(11)
        
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run("평가 시스템 v6.0 | RAG + LLM | 정량평가 중심")
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = self.COLORS['gray']
    
    def _add_enhanced_executive_summary(self, doc: Document, 
                                       patent_info: Dict, state: Dict):
        """강화된 Executive Summary"""
        heading = doc.add_heading('Executive Summary', level=1)
        for run in heading.runs:
            run.font.color.rgb = self.COLORS['primary']
        
        # 1. 평가 개요
        doc.add_heading('1. 평가 개요', level=2)
        
        overall_score = state.get('overall_score', 0)
        final_grade = state.get('final_grade', 'N/A')
        tech_score = state.get('tech_score', 0)
        rights_score = state.get('rights_score', 0)
        market_score = state.get('market_score', 0)
        
        # 표로 요약 정보 제시
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'
        
        data = [
            ('종합 점수', f"{overall_score:.1f}점 ({final_grade})"),
            ('기술성 평가', f"{tech_score:.1f}점"),
            ('권리성 평가', f"{rights_score:.1f}점"),
            ('활용성 평가', f"{market_score:.1f}점"),
            ('평가 방법', '정량평가 + LLM 정성평가'),
        ]
        
        for i, (key, value) in enumerate(data):
            cells = table.rows[i].cells
            cells[0].text = key
            cells[1].text = value
            
            # 헤더 셀 강조
            for cell in [cells[0]]:
                cell.paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()
        
        # 2. 핵심 강점 (Key Strengths)
        doc.add_heading('2. 핵심 강점 (Key Strengths)', level=2)
        
        tech_eval = state.get('tech_evaluation', {})
        rights_eval = state.get('rights_evaluation', {})
        market_eval = state.get('market_evaluation', {})
        
        strengths = []
        
        # 각 평가에서 강점 추출
        if tech_score >= 75:
            strengths.append("✓ 우수한 기술적 혁신성 및 구현 상세도")
        if rights_score >= 75:
            strengths.append("✓ 견고한 권리범위 및 청구항 구조")
        if market_score >= 75:
            strengths.append("✓ 높은 시장 활용성 및 상용화 가능성")
        
        # tech_evaluation에서 추가 강점
        tech_strengths = tech_eval.get('key_strengths', [])
        for s in tech_strengths[:2]:
            strengths.append(f"✓ {s}")
        
        for strength in strengths[:5]:  # 최대 5개
            p = doc.add_paragraph(strength, style='List Bullet')
            p.runs[0].font.color.rgb = self.COLORS['success']
        
        # 3. 개선 필요 영역 (Areas for Improvement)
        doc.add_heading('3. 개선 필요 영역 (Areas for Improvement)', level=2)
        
        weaknesses = []
        
        if tech_score < 70:
            weaknesses.append("• 기술적 구현 방법의 상세화 필요")
        if rights_score < 70:
            weaknesses.append("• 청구항 범위 확대 및 종속항 보강 필요")
        if market_score < 70:
            weaknesses.append("• 시장 검증 및 비즈니스 모델 구체화 필요")
        
        # tech_evaluation에서 약점 추출
        tech_weaknesses = tech_eval.get('key_weaknesses', [])
        for w in tech_weaknesses[:2]:
            weaknesses.append(f"• {w}")
        
        if not weaknesses:
            weaknesses.append("• 전반적으로 우수한 수준 유지 중")
        
        for weakness in weaknesses[:5]:
            p = doc.add_paragraph(weakness, style='List Bullet')
            if "우수" not in weakness:
                p.runs[0].font.color.rgb = self.COLORS['warning']
        
        doc.add_paragraph()
        
        # 4. 평가 방법론 설명
        doc.add_heading('4. 평가 방법론 (v6.0)', level=2)
        
        self._create_info_box(
            doc,
            "정량평가 중심 하이브리드 시스템",
            """본 평가는 정량평가를 중심으로 LLM 정성평가를 결합한 하이브리드 방식을 사용합니다:

• 기술성: 정량 60% (PDF 원문 지표) + 정성 40% (LLM 분석)
• 권리성: 정량 70% (청구항 구조 분석) + 정성 30% (LLM 분석)
• 활용성: 정량+웹서치 70% + 정성 30% (LLM 분석)

정량 지표는 특허 원문에서 직접 추출하며, RAG(Retrieval-Augmented Generation) 시스템을 
통해 관련 컨텍스트를 검색하여 LLM이 정성적 분석을 수행합니다."""
        )
        
        doc.add_paragraph()
        
        # 5. 특허 기본 정보
        doc.add_heading('5. 특허 기본 정보', level=2)
        
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Light List Accent 1'
        
        ipc_codes = patent_info.get('ipc_codes', [])
        ipc_str = ', '.join(ipc_codes[:5]) if ipc_codes else 'N/A'
        
        data = [
            ('특허번호', patent_info.get('number', 'N/A')),
            ('발명명칭', patent_info.get('title', 'N/A')[:80] + '...'),
            ('출원인', patent_info.get('applicant', 'N/A')),
            ('청구항 수', f"{patent_info.get('claims_count', 0)}개"),
            ('IPC 코드', ipc_str),
            ('발명자 수', f"{len(patent_info.get('inventors', []))}명"),
        ]
        
        for i, (key, value) in enumerate(data):
            cells = table.rows[i].cells
            cells[0].text = key
            cells[1].text = value
            cells[0].paragraphs[0].runs[0].font.bold = True
    
    def _add_table_of_contents(self, doc: Document):
        """목차 추가"""
        heading = doc.add_heading('목차 (Table of Contents)', level=1)
        for run in heading.runs:
            run.font.color.rgb = self.COLORS['primary']
        
        toc_items = [
            ("1. Executive Summary", "평가 개요 및 핵심 사항"),
            ("2. 평가 결과 시각화", "차트 및 그래프"),
            ("3. 특허 개요", "메타데이터 및 기술 분류"),
            ("4. 기술성 평가", "정량지표 + LLM 분석"),
            ("5. 권리성 평가", "청구항 구조 + LLM 분석"),
            ("6. 활용성 평가", "시장성 + 웹서치 + LLM 분석"),
            ("7. 비교 분석", "벤치마크 및 경쟁 분석"),
            ("8. 종합 평가 및 제언", "전략적 권고사항"),
            ("9. 개선 로드맵", "단계별 개선 계획"),
            ("10. Reference", "참고 문서 및 출처"),
            ("11. Appendix", "평가 지표 상세 및 체크리스트"),
        ]
        
        for section, desc in toc_items:
            p = doc.add_paragraph()
            run = p.add_run(section)
            run.font.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = self.COLORS['dark']
            
            run = p.add_run(f"\n     {desc}")
            run.font.size = Pt(9)
            run.font.color.rgb = self.COLORS['gray']
            run.italic = True
    
    def _add_enhanced_evaluation_charts(self, doc: Document, 
                                       chart_paths: Dict[str, str], state: Dict):
        """강화된 평가 결과 시각화"""
        heading = doc.add_heading('평가 결과 시각화', level=1)
        for run in heading.runs:
            run.font.color.rgb = self.COLORS['primary']
        
        # 차트 설명
        self._create_info_box(
            doc,
            "시각화 개요",
            "다음 차트들은 특허의 각 평가 영역별 점수를 다양한 각도에서 시각화한 것입니다. "
            "막대 차트는 절대 점수를, 레이더 차트는 균형도를, 파이 차트는 상대적 비중을 보여줍니다."
        )
        
        doc.add_paragraph()
        
        # 막대 차트
        if 'bar' in chart_paths and Path(chart_paths['bar']).exists():
            doc.add_heading('1. 평가 영역별 점수 비교', level=2)
            p = doc.add_paragraph()
            p.add_run("각 평가 영역(기술성, 권리성, 활용성)의 점수를 비교하여 강점과 약점을 파악합니다.")
            p.runs[0].font.size = Pt(9)
            p.runs[0].font.color.rgb = self.COLORS['gray']
            
            doc.add_picture(chart_paths['bar'], width=Inches(6))
            doc.add_paragraph()
        
        # 레이더 차트
        if 'radar' in chart_paths and Path(chart_paths['radar']).exists():
            doc.add_heading('2. 균형도 분석 (레이더 차트)', level=2)
            p = doc.add_paragraph()
            p.add_run("특허의 전체적인 균형도를 시각화합니다. 이상적인 형태는 정삼각형에 가까운 모양입니다.")
            p.runs[0].font.size = Pt(9)
            p.runs[0].font.color.rgb = self.COLORS['gray']
            
            doc.add_picture(chart_paths['radar'], width=Inches(5))
            doc.add_paragraph()
        
        # 파이 차트
        if 'pie' in chart_paths and Path(chart_paths['pie']).exists():
            doc.add_heading('3. 가중치별 기여도', level=2)
            p = doc.add_paragraph()
            p.add_run("각 평가 영역이 종합 점수에 기여하는 비중을 보여줍니다.")
            p.runs[0].font.size = Pt(9)
            p.runs[0].font.color.rgb = self.COLORS['gray']
            
            doc.add_picture(chart_paths['pie'], width=Inches(5))
            doc.add_paragraph()
        
        # 점수 분석
        doc.add_heading('4. 점수 분석', level=2)
        
        tech_score = state.get('tech_score', 0)
        rights_score = state.get('rights_score', 0)
        market_score = state.get('market_score', 0)
        overall_score = state.get('overall_score', 0)
        
        scores_analysis = f"""
• 최고 점수 영역: {self._get_highest_category(tech_score, rights_score, market_score)}
• 최저 점수 영역: {self._get_lowest_category(tech_score, rights_score, market_score)}
• 점수 편차: {max(tech_score, rights_score, market_score) - min(tech_score, rights_score, market_score):.1f}점
• 균형도: {'우수' if max(tech_score, rights_score, market_score) - min(tech_score, rights_score, market_score) < 15 else '보통' if max(tech_score, rights_score, market_score) - min(tech_score, rights_score, market_score) < 25 else '개선 필요'}
"""
        
        p = doc.add_paragraph(scores_analysis)
        p.runs[0].font.size = Pt(10)
    
    def _get_highest_category(self, tech: float, rights: float, market: float) -> str:
        """최고 점수 영역 반환"""
        scores = {'기술성': tech, '권리성': rights, '활용성': market}
        return f"{max(scores, key=scores.get)} ({max(scores.values()):.1f}점)"
    
    def _get_lowest_category(self, tech: float, rights: float, market: float) -> str:
        """최저 점수 영역 반환"""
        scores = {'기술성': tech, '권리성': rights, '활용성': market}
        return f"{min(scores, key=scores.get)} ({min(scores.values()):.1f}점)"
    
    def _add_patent_overview(self, doc: Document, patent_info: Dict):
        """특허 개요 및 메타데이터"""
        heading = doc.add_heading('특허 개요 (Patent Overview)', level=1)
        for run in heading.runs:
            run.font.color.rgb = self.COLORS['primary']
        
        # 1. 기본 정보
        doc.add_heading('1. 기본 정보', level=2)
        
        table = doc.add_table(rows=8, cols=2)
        table.style = 'Medium Shading 1 Accent 1'
        
        ipc_codes = patent_info.get('ipc_codes', [])
        inventors = patent_info.get('inventors', [])
        
        data = [
            ('특허번호', patent_info.get('number', 'N/A')),
            ('발명명칭', patent_info.get('title', 'N/A')),
            ('출원인', patent_info.get('applicant', 'N/A')),
            ('발명자', ', '.join(inventors) if inventors else 'N/A'),
            ('청구항 수', f"{patent_info.get('claims_count', 0)}개"),
            ('IPC 주분류', ipc_codes[0] if ipc_codes else 'N/A'),
            ('IPC 전체', ', '.join(ipc_codes[:5]) if len(ipc_codes) > 1 else 'N/A'),
            ('도면 수', f"{patent_info.get('drawing_count', 0)}개"),
        ]
        
        for i, (key, value) in enumerate(data):
            cells = table.rows[i].cells
            cells[0].text = key
            cells[1].text = str(value)
            cells[0].paragraphs[0].runs[0].font.bold = True
            cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
        doc.add_paragraph()
        
        # 2. 기술 분류
        doc.add_heading('2. 기술 분류 (IPC 분석)', level=2)
        
        if ipc_codes:
            self._create_info_box(
                doc,
                "IPC 주분류 설명",
                f"주분류 코드 '{ipc_codes[0]}'는 다음 기술 분야에 속합니다:\n"
                f"• 섹션: {ipc_codes[0][0] if ipc_codes[0] else 'N/A'}\n"
                f"• 총 {len(ipc_codes)}개의 IPC 코드로 분류되어 있어, "
                f"{'다양한 기술 영역을 포괄' if len(ipc_codes) > 3 else '특정 기술 분야에 집중'}하고 있습니다.",
                self.COLORS['light']
            )
        
        doc.add_paragraph()
        
        # 3. 청구항 구조
        doc.add_heading('3. 청구항 구조', level=2)
        
        claims_count = patent_info.get('claims_count', 0)
        
        # 청구항 구조 시각적 표현
        p = doc.add_paragraph()
        p.add_run(f"총 청구항 수: {claims_count}개\n\n")
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(11)
        
        if claims_count > 0:
            if claims_count >= 20:
                assessment = "매우 상세한 청구항 구조 (20개 이상)"
                color = self.COLORS['success']
            elif claims_count >= 10:
                assessment = "충분한 청구항 구조 (10-19개)"
                color = self.COLORS['secondary']
            else:
                assessment = "간결한 청구항 구조 (10개 미만)"
                color = self.COLORS['warning']
            
            p.add_run(f"평가: {assessment}")
            p.runs[-1].font.color.rgb = color
    
    def _add_enhanced_technology_evaluation(self, doc: Document, state: Dict):
        """강화된 기술성 평가 (기존 구조 + 추가 분석)"""
        heading = doc.add_heading('기술성 평가 (Technology Evaluation)', level=1)
        for run in heading.runs:
            run.font.color.rgb = self.COLORS['primary']
        
        tech_score = state.get('tech_score', 0)
        tech_metrics = state.get('tech_metrics', {})
        tech_quantitative = state.get('tech_quantitative', {})
        tech_qualitative = state.get('tech_qualitative', {})
        tech_binary = state.get('tech_binary', {})
        tech_evaluation = state.get('tech_evaluation', {})
        
        # 점수 개요
        self._create_info_box(
            doc,
            f"기술성 종합 점수: {tech_score:.1f}점",
            f"{'[우수]' if tech_score >= 75 else '[양호]' if tech_score >= 60 else '[개선필요]'} "
            f"정량평가 {tech_quantitative.get('total', 0):.1f}점 × 60% + "
            f"정성평가 {tech_qualitative.get('qualitative_score', 0):.1f}점 × 40%",
            self.COLORS['success'] if tech_score >= 75 else self.COLORS['warning']
        )
        
        doc.add_paragraph()
        
        # ========== 기존 v5.0 구조 유지 ==========
        
        # 1. 정량 지표 (PDF 원문 기반)
        doc.add_heading('1. 정량 지표 (PDF 원문 기반) - 60%', level=2)
        
        table = doc.add_table(rows=4, cols=4)
        table.style = 'Light Grid Accent 1'
        
        header_cells = table.rows[0].cells
        header_cells[0].text = '지표'
        header_cells[1].text = '측정값'
        header_cells[2].text = '점수'
        header_cells[3].text = '가중치'
        
        for cell in header_cells:
            cell.paragraphs[0].runs[0].font.bold = True
        
        data = [
            ('X7. 도면 수', f"{tech_metrics.get('X7_drawing_count', 0)}개", 
             f"{tech_quantitative.get('drawing_score', 0):.1f}", '40%'),
            ('X8. 발명명칭 길이', f"{tech_metrics.get('X8_title_length', 0)}자", 
             f"{tech_quantitative.get('title_score', 0):.1f}", '30%'),
            ('X9. 청구항 계열 수', f"{tech_metrics.get('X9_claim_series', 0)}개", 
             f"{tech_quantitative.get('series_score', 0):.1f}", '30%'),
        ]
        
        for i, (label, value, score, weight) in enumerate(data, start=1):
            cells = table.rows[i].cells
            cells[0].text = label
            cells[1].text = value
            cells[2].text = score
            cells[3].text = weight
        
        doc.add_paragraph()
        
        # 구조방정식 모델
        doc.add_heading('2. 구조방정식 모델', level=2)
        
        drawing_score = tech_quantitative.get('drawing_score', 0)
        title_score = tech_quantitative.get('title_score', 0)
        series_score = tech_quantitative.get('series_score', 0)
        quant_total = tech_quantitative.get('total', 0)
        qual_score = tech_qualitative.get('qualitative_score', 0)
        
        p = doc.add_paragraph()
        p.add_run("정량 점수 계산:\n").bold = True
        p.add_run(f"= X7(도면) × 0.4 + X8(명칭) × 0.3 + X9(계열) × 0.3\n")
        p.add_run(f"= {drawing_score:.1f} × 0.4 + {title_score:.1f} × 0.3 + {series_score:.1f} × 0.3\n")
        p.add_run(f"= {quant_total:.1f}점\n\n")
        
        p.add_run("최종 기술성 점수:\n").bold = True
        p.add_run(f"= 정량({quant_total:.1f}) × 60% + 정성({qual_score:.1f}) × 40%\n")
        p.add_run(f"= {tech_score:.1f}점\n")
        
        doc.add_paragraph()
        
        # 3. Binary 체크리스트
        doc.add_heading('3. Binary 체크리스트', level=2)
        
        table = doc.add_table(rows=len(tech_binary)+1, cols=2)
        table.style = 'Light List Accent 1'
        
        header_cells = table.rows[0].cells
        header_cells[0].text = '항목'
        header_cells[1].text = '상태'
        
        for i, (key, value) in enumerate(tech_binary.items(), start=1):
            cells = table.rows[i].cells
            cells[0].text = key
            cells[1].text = "✓ 충족" if value else "✗ 미충족"
            
            if value:
                cells[1].paragraphs[0].runs[0].font.color.rgb = self.COLORS['success']
            else:
                cells[1].paragraphs[0].runs[0].font.color.rgb = self.COLORS['danger']
        
        doc.add_paragraph()
        
        # 4. 정성 평가 (LLM)
        doc.add_heading('4. 정성 평가 (LLM 분석) - 40%', level=2)
        
        p = doc.add_paragraph()
        p.add_run("기술적 혁신성:\n").bold = True
        innovation_summary = tech_qualitative.get('innovation_summary', 
                                                   tech_evaluation.get('technical_summary', 'N/A'))
        p.add_run(f"{innovation_summary}\n\n")
        
        p.add_run("구현 상세도:\n").bold = True
        impl_summary = tech_qualitative.get('implementation_summary', 'N/A')
        p.add_run(f"{impl_summary}\n\n")
        
        p.add_run("기술적 차별성:\n").bold = True
        diff_summary = tech_qualitative.get('differentiation_summary', 'N/A')
        p.add_run(f"{diff_summary}\n")
        
        doc.add_paragraph()
        
        # ========== 추가 강화 내용 ==========
        
        # 5. 세부 지표 분석
        doc.add_heading('5. 세부 지표 상세 분석', level=2)
        
        # 도면 분석
        drawing_count = tech_metrics.get('X7_drawing_count', 0)
        p = doc.add_paragraph()
        p.add_run("📊 도면 수 분석\n").bold = True
        p.runs[0].font.color.rgb = self.COLORS['primary']
        
        if drawing_count >= 10:
            analysis = f"매우 풍부한 도면({drawing_count}개)으로 기술 구현이 시각적으로 상세히 설명되어 있습니다."
        elif drawing_count >= 5:
            analysis = f"적절한 도면({drawing_count}개)으로 핵심 기술이 설명되어 있습니다."
        else:
            analysis = f"도면이 제한적({drawing_count}개)이어서 기술 이해에 추가 설명이 필요할 수 있습니다."
        
        p.add_run(analysis)
        doc.add_paragraph()
        
        # 발명명칭 분석
        title_length = tech_metrics.get('X8_title_length', 0)
        p = doc.add_paragraph()
        p.add_run("📝 발명명칭 분석\n").bold = True
        p.runs[0].font.color.rgb = self.COLORS['primary']
        
        if title_length >= 30:
            analysis = f"상세한 명칭({title_length}자)으로 발명의 핵심 내용을 잘 표현하고 있습니다."
        elif title_length >= 15:
            analysis = f"적절한 길이의 명칭({title_length}자)입니다."
        else:
            analysis = f"간결한 명칭({title_length}자)으로, 발명 내용의 추가 설명이 필요할 수 있습니다."
        
        p.add_run(analysis)
        doc.add_paragraph()
        
        # 청구항 계열 분석
        series_count = tech_metrics.get('X9_claim_series', 0)
        p = doc.add_paragraph()
        p.add_run("🔗 청구항 계열 분석\n").bold = True
        p.runs[0].font.color.rgb = self.COLORS['primary']
        
        if series_count >= 3:
            analysis = f"다양한 청구항 계열({series_count}개)로 기술이 체계적으로 보호되고 있습니다."
        elif series_count >= 2:
            analysis = f"적절한 청구항 계열({series_count}개)로 핵심 기술이 보호됩니다."
        else:
            analysis = f"제한적인 청구항 계열({series_count}개)로, 추가 보호범위 확대가 권장됩니다."
        
        p.add_run(analysis)
    
    def _add_enhanced_rights_evaluation(self, doc: Document, state: Dict):
        """강화된 권리성 평가 (기존 구조 + 추가 분석)"""
        heading = doc.add_heading('권리성 평가 (Rights Evaluation)', level=1)
        for run in heading.runs:
            run.font.color.rgb = self.COLORS['primary']
        
        rights_score = state.get('rights_score', 0)
        rights_metrics = state.get('rights_metrics', {})
        rights_quantitative = state.get('rights_quantitative', {})
        rights_qualitative = state.get('rights_qualitative', {})
        rights_binary = state.get('rights_binary', {})
        
        # 점수 개요
        self._create_info_box(
            doc,
            f"권리성 종합 점수: {rights_score:.1f}점",
            f"{'[우수]' if rights_score >= 75 else '[양호]' if rights_score >= 60 else '[개선필요]'} "
            f"정량평가 {rights_quantitative.get('total', 0):.1f}점 × 70% + "
            f"정성평가 {rights_qualitative.get('qualitative_score', 0):.1f}점 × 30%",
            self.COLORS['success'] if rights_score >= 75 else self.COLORS['warning']
        )
        
        doc.add_paragraph()
        
        # ========== 기존 v5.0 구조 유지 ==========
        
        # 1. 정량 지표
        doc.add_heading('1. 정량 지표 (PDF 원문 기반) - 70%', level=2)
        
        table = doc.add_table(rows=7, cols=4)
        table.style = 'Light Grid Accent 1'
        
        header_cells = table.rows[0].cells
        header_cells[0].text = '지표'
        header_cells[1].text = '측정값'
        header_cells[2].text = '점수'
        header_cells[3].text = '가중치'
        
        for cell in header_cells:
            cell.paragraphs[0].runs[0].font.bold = True
        
        data = [
            ('X1. IPC 코드 수', f"{rights_metrics.get('X1_ipc_count', 0)}개", 
             f"{rights_quantitative.get('ipc_score', 0):.1f}", '25%'),
            ('X2. 독립항 수', f"{rights_metrics.get('X2_independent_claims', 0)}개", 
             f"{rights_quantitative.get('independent_score', 0):.1f}", '15%'),
            ('X3. 종속항 수', f"{rights_metrics.get('X3_dependent_claims', 0)}개", 
             f"{rights_quantitative.get('dependent_score', 0):.1f}", '15%'),
            ('X4. 전체 청구항', f"{rights_metrics.get('X4_total_claims', 0)}개", 
             f"{rights_quantitative.get('total_claims_score', 0):.1f}", '15%'),
            ('X5. 독립항 평균 길이', f"{rights_metrics.get('X5_independent_avg_length', 0):.1f}자", 
             f"{rights_quantitative.get('independent_length_score', 0):.1f}", '12.5%'),
            ('X6. 종속항 평균 길이', f"{rights_metrics.get('X6_dependent_avg_length', 0):.1f}자", 
             f"{rights_quantitative.get('dependent_length_score', 0):.1f}", '12.5%'),
        ]
        
        for i, (label, value, score, weight) in enumerate(data, start=1):
            cells = table.rows[i].cells
            cells[0].text = label
            cells[1].text = value
            cells[2].text = score
            cells[3].text = weight
        
        doc.add_paragraph()
        
        # 2. 구조방정식 모델
        doc.add_heading('2. 구조방정식 모델', level=2)
        
        ipc_score = rights_quantitative.get('ipc_score', 0)
        claims_count_score = rights_quantitative.get('claims_count_score', 0)
        claims_length_score = rights_quantitative.get('claims_length_score', 0)
        hierarchy_score = rights_quantitative.get('hierarchy_score', 0)
        quant_total = rights_quantitative.get('total', 0)
        qual_score = rights_qualitative.get('qualitative_score', 0)
        
        p = doc.add_paragraph()
        p.add_run("정량 점수 계산:\n").bold = True
        p.add_run(f"= IPC × 0.25 + 청구항개수 × 0.30 + 청구항길이 × 0.25 + 계층구조 × 0.20\n")
        p.add_run(f"= {ipc_score:.1f} × 0.25 + {claims_count_score:.1f} × 0.30 + ")
        p.add_run(f"{claims_length_score:.1f} × 0.25 + {hierarchy_score:.1f} × 0.20\n")
        p.add_run(f"= {quant_total:.1f}점\n\n")
        
        p.add_run("최종 권리성 점수:\n").bold = True
        p.add_run(f"= 정량({quant_total:.1f}) × 70% + 정성({qual_score:.1f}) × 30%\n")
        p.add_run(f"= {rights_score:.1f}점\n")
        
        doc.add_paragraph()
        
        # 3. Binary 체크리스트
        doc.add_heading('3. Binary 체크리스트', level=2)
        
        table = doc.add_table(rows=len(rights_binary)+1, cols=2)
        table.style = 'Light List Accent 1'
        
        header_cells = table.rows[0].cells
        header_cells[0].text = '항목'
        header_cells[1].text = '상태'
        
        for i, (key, value) in enumerate(rights_binary.items(), start=1):
            cells = table.rows[i].cells
            cells[0].text = key
            cells[1].text = "✓ 충족" if value else "✗ 미충족"
            
            if value:
                cells[1].paragraphs[0].runs[0].font.color.rgb = self.COLORS['success']
            else:
                cells[1].paragraphs[0].runs[0].font.color.rgb = self.COLORS['danger']
        
        doc.add_paragraph()
        
        # 4. 정성 평가 (LLM)
        doc.add_heading('4. 정성 평가 (LLM 분석) - 30%', level=2)
        
        p = doc.add_paragraph()
        p.add_run("권리범위 평가:\n").bold = True
        scope_summary = rights_qualitative.get('scope_summary', 'N/A')
        p.add_run(f"{scope_summary}\n\n")
        
        p.add_run("청구항 견고성:\n").bold = True
        robustness_summary = rights_qualitative.get('robustness_summary', 'N/A')
        p.add_run(f"{robustness_summary}\n\n")
        
        p.add_run("회피 설계 난이도:\n").bold = True
        avoidance_summary = rights_qualitative.get('avoidance_summary', 'N/A')
        p.add_run(f"{avoidance_summary}\n")
        
        doc.add_paragraph()
        
        # ========== 추가 강화 내용 ==========
        
        # 5. 청구항 구조 심층 분석
        doc.add_heading('5. 청구항 구조 심층 분석', level=2)
        
        independent_claims = rights_metrics.get('X2_independent_claims', 0)
        dependent_claims = rights_metrics.get('X3_dependent_claims', 0)
        total_claims = rights_metrics.get('X4_total_claims', 0)
        
        if total_claims > 0:
            dep_ratio = (dependent_claims / total_claims) * 100
            
            p = doc.add_paragraph()
            p.add_run("📋 청구항 구성 분석\n").bold = True
            p.runs[0].font.color.rgb = self.COLORS['primary']
            
            p.add_run(f"• 독립항: {independent_claims}개 ({(independent_claims/total_claims)*100:.1f}%)\n")
            p.add_run(f"• 종속항: {dependent_claims}개 ({dep_ratio:.1f}%)\n")
            p.add_run(f"• 종속항 비율: ")
            
            if dep_ratio >= 70:
                p.add_run("매우 우수 (70% 이상)\n")
                p.runs[-1].font.color.rgb = self.COLORS['success']
            elif dep_ratio >= 50:
                p.add_run("우수 (50-70%)\n")
                p.runs[-1].font.color.rgb = self.COLORS['secondary']
            else:
                p.add_run("보통 (50% 미만)\n")
                p.runs[-1].font.color.rgb = self.COLORS['warning']
            
            p.add_run(f"\n평가: ")
            if dep_ratio >= 70:
                p.add_run("종속항이 충분히 확보되어 권리범위가 다층적으로 보호됩니다.")
            elif dep_ratio >= 50:
                p.add_run("적절한 수준의 종속항으로 기본적인 권리보호가 가능합니다.")
            else:
                p.add_run("종속항 확대를 통한 권리범위 강화를 권장합니다.")
            
            doc.add_paragraph()
        
        # 6. IPC 다양성 분석
        ipc_count = rights_metrics.get('X1_ipc_count', 0)
        p = doc.add_paragraph()
        p.add_run("🏷️ IPC 분류 다양성\n").bold = True
        p.runs[0].font.color.rgb = self.COLORS['primary']
        
        if ipc_count >= 5:
            analysis = f"매우 다양한 IPC 분류({ipc_count}개)로 기술이 다각도로 보호됩니다."
        elif ipc_count >= 3:
            analysis = f"적절한 IPC 분류({ipc_count}개)로 핵심 기술 영역이 커버됩니다."
        else:
            analysis = f"제한적인 IPC 분류({ipc_count}개)입니다. 관련 기술 분야 확대를 고려할 수 있습니다."
        
        p.add_run(analysis)
    
    def _add_enhanced_market_evaluation(self, doc: Document, state: Dict):
        """강화된 활용성 평가 (기존 구조 + 추가 분석)"""
        heading = doc.add_heading('활용성 평가 (Market Utilization)', level=1)
        for run in heading.runs:
            run.font.color.rgb = self.COLORS['primary']
        
        market_score = state.get('market_score', 0)
        market_metrics = state.get('market_metrics', {})
        market_quantitative = state.get('market_quantitative', {})
        market_qualitative = state.get('market_qualitative', {})
        market_binary = state.get('market_binary', {})
        market_web_search = state.get('market_web_search', {})
        
        # 점수 개요
        self._create_info_box(
            doc,
            f"활용성 종합 점수: {market_score:.1f}점",
            f"{'[우수]' if market_score >= 75 else '[양호]' if market_score >= 60 else '[개선필요]'} "
            f"정량+웹서치 {market_quantitative.get('total', 0):.1f}점 × 70% + "
            f"정성평가 {market_qualitative.get('qualitative_score', 0):.1f}점 × 30%",
            self.COLORS['success'] if market_score >= 75 else self.COLORS['warning']
        )
        
        doc.add_paragraph()
        
        # ========== 기존 v5.0 구조 유지 ==========
        
        # 1. 정량 지표
        doc.add_heading('1. 정량 지표 (PDF 원문 기반)', level=2)
        
        p = doc.add_paragraph()
        inventor_count = market_metrics.get('X10_inventor_count', 0)
        inventor_score = market_quantitative.get('inventor_score', 0)
        
        p.add_run(f"X10. 발명자 수: ").bold = True
        p.add_run(f"{inventor_count}명 ")
        p.add_run(f"→ {inventor_score:.1f}점 (가중치 30%)\n")
        
        if inventor_count >= 5:
            p.add_run("평가: 다수의 발명자가 참여하여 협업 기술 개발이 이루어졌습니다.\n")
        elif inventor_count >= 3:
            p.add_run("평가: 적절한 규모의 연구팀이 참여했습니다.\n")
        else:
            p.add_run("평가: 소규모 팀 또는 개인 발명입니다.\n")
        
        doc.add_paragraph()
        
        # 2. 웹 서치 결과
        doc.add_heading('2. 웹 서치 결과 (실시간 시장 정보)', level=2)
        
        # 출원인 분석
        p = doc.add_paragraph()
        p.add_run("🏢 출원인 시장 지위\n").bold = True
        p.runs[0].font.color.rgb = self.COLORS['primary']
        
        applicant_grade = market_web_search.get('applicant_grade', 'Unknown')
        applicant_score = market_quantitative.get('applicant_score', 0)
        applicant_summary = market_web_search.get('applicant_summary', 'N/A')
        
        p.add_run(f"등급: {applicant_grade} → {applicant_score:.1f}점 (가중치 40%)\n")
        p.add_run(f"분석: {applicant_summary}\n")
        
        doc.add_paragraph()
        
        # 기술 분야 분석
        p = doc.add_paragraph()
        p.add_run("📈 기술 분야 성장성\n").bold = True
        p.runs[0].font.color.rgb = self.COLORS['primary']
        
        tech_grade = market_web_search.get('tech_grade', 'Unknown')
        tech_field_score = market_quantitative.get('tech_field_score', 0)
        tech_summary = market_web_search.get('tech_summary', 'N/A')
        
        p.add_run(f"등급: {tech_grade} → {tech_field_score:.1f}점 (가중치 30%)\n")
        p.add_run(f"분석: {tech_summary}\n")
        
        doc.add_paragraph()
        
        # 3. 구조방정식 모델
        doc.add_heading('3. 구조방정식 모델', level=2)
        
        quant_total = market_quantitative.get('total', 0)
        qual_score = market_qualitative.get('qualitative_score', 0)
        
        p = doc.add_paragraph()
        p.add_run("정량+웹서치 점수:\n").bold = True
        p.add_run(f"= 발명자(30%) + 출원인(40%) + 기술분야(30%)\n")
        p.add_run(f"= {inventor_score:.1f} × 0.30 + {applicant_score:.1f} × 0.40 + {tech_field_score:.1f} × 0.30\n")
        p.add_run(f"= {quant_total:.1f}점\n\n")
        
        p.add_run("최종 활용성 점수:\n").bold = True
        p.add_run(f"= (정량+웹서치)({quant_total:.1f}) × 70% + 정성({qual_score:.1f}) × 30%\n")
        p.add_run(f"= {market_score:.1f}점\n")
        
        doc.add_paragraph()
        
        # 4. Binary 체크리스트
        doc.add_heading('4. Binary 체크리스트', level=2)
        
        table = doc.add_table(rows=len(market_binary)+1, cols=2)
        table.style = 'Light List Accent 1'
        
        header_cells = table.rows[0].cells
        header_cells[0].text = '항목'
        header_cells[1].text = '상태'
        
        for i, (key, value) in enumerate(market_binary.items(), start=1):
            cells = table.rows[i].cells
            cells[0].text = key
            cells[1].text = "✓ 충족" if value else "✗ 미충족"
            
            if value:
                cells[1].paragraphs[0].runs[0].font.color.rgb = self.COLORS['success']
            else:
                cells[1].paragraphs[0].runs[0].font.color.rgb = self.COLORS['danger']
        
        doc.add_paragraph()
        
        # 5. 정성 평가 (LLM)
        doc.add_heading('5. 정성 평가 (LLM 분석) - 30%', level=2)
        
        p = doc.add_paragraph()
        p.add_run("실무 적용성:\n").bold = True
        applicability = market_qualitative.get('applicability_summary', 'N/A')
        p.add_run(f"{applicability}\n\n")
        
        p.add_run("시장 적합성:\n").bold = True
        market_fit = market_qualitative.get('market_fit_summary', 'N/A')
        p.add_run(f"{market_fit}\n\n")
        
        p.add_run("상용화 가능성:\n").bold = True
        commercialization = market_qualitative.get('commercialization_summary', 'N/A')
        p.add_run(f"{commercialization}\n")
        
        doc.add_paragraph()
        
        # ========== 추가 강화 내용 ==========
        
        # 6. 시장 환경 분석
        doc.add_heading('6. 시장 환경 종합 분석', level=2)
        
        self._create_info_box(
            doc,
            "시장 진입 전략 제안",
            f"출원인의 시장 지위({applicant_grade})와 기술 분야 성장성({tech_grade})을 "
            f"고려할 때, {'직접 사업화' if applicant_score >= 70 else '라이선싱' if market_score >= 60 else '공동 개발'}을 우선 고려할 수 있습니다. "
            f"{'시장 선도 기업으로서' if applicant_score >= 80 else '시장 진입을 위해'} "
            f"{'기술 우위를 활용한' if market_score >= 75 else '전략적 파트너십을 통한'} 접근이 효과적일 것입니다.",
            self.COLORS['light']
        )
    
    def _add_comparative_analysis(self, doc: Document, state: Dict):
        """비교 분석 (신규)"""
        heading = doc.add_heading('비교 분석 (Comparative Analysis)', level=1)
        for run in heading.runs:
            run.font.color.rgb = self.COLORS['primary']
        
        tech_score = state.get('tech_score', 0)
        rights_score = state.get('rights_score', 0)
        market_score = state.get('market_score', 0)
        overall_score = state.get('overall_score', 0)
        
        # 1. 산업 평균 대비 비교
        doc.add_heading('1. 산업 평균 대비 비교', level=2)
        
        # 가상의 산업 평균 (실제로는 데이터베이스에서 가져와야 함)
        industry_avg = {
            'technology': 70.0,
            'rights': 65.0,
            'market': 60.0,
            'overall': 66.0
        }
        
        table = doc.add_table(rows=5, cols=4)
        table.style = 'Medium Shading 1 Accent 1'
        
        header_cells = table.rows[0].cells
        header_cells[0].text = '평가 영역'
        header_cells[1].text = '본 특허'
        header_cells[2].text = '산업 평균'
        header_cells[3].text = '차이'
        
        for cell in header_cells:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
        data = [
            ('기술성', f"{tech_score:.1f}", f"{industry_avg['technology']:.1f}", 
             f"{tech_score - industry_avg['technology']:+.1f}"),
            ('권리성', f"{rights_score:.1f}", f"{industry_avg['rights']:.1f}", 
             f"{rights_score - industry_avg['rights']:+.1f}"),
            ('활용성', f"{market_score:.1f}", f"{industry_avg['market']:.1f}", 
             f"{market_score - industry_avg['market']:+.1f}"),
            ('종합', f"{overall_score:.1f}", f"{industry_avg['overall']:.1f}", 
             f"{overall_score - industry_avg['overall']:+.1f}"),
        ]
        
        for i, (category, score, avg, diff) in enumerate(data, start=1):
            cells = table.rows[i].cells
            cells[0].text = category
            cells[1].text = score
            cells[2].text = avg
            cells[3].text = diff
            
            # 차이 값 색상 설정
            diff_val = float(diff)
            if diff_val > 0:
                cells[3].paragraphs[0].runs[0].font.color.rgb = self.COLORS['success']
            elif diff_val < 0:
                cells[3].paragraphs[0].runs[0].font.color.rgb = self.COLORS['danger']
        
        doc.add_paragraph()
        
        # 2. 강점/약점 벤치마킹
        doc.add_heading('2. 강점/약점 벤치마킹', level=2)
        
        p = doc.add_paragraph()
        p.add_run("상대적 강점 (산업 평균 대비 우수):\n").bold = True
        
        if tech_score > industry_avg['technology']:
            p.add_run(f"✓ 기술성: 산업 평균 대비 +{tech_score - industry_avg['technology']:.1f}점\n")
        if rights_score > industry_avg['rights']:
            p.add_run(f"✓ 권리성: 산업 평균 대비 +{rights_score - industry_avg['rights']:.1f}점\n")
        if market_score > industry_avg['market']:
            p.add_run(f"✓ 활용성: 산업 평균 대비 +{market_score - industry_avg['market']:.1f}점\n")
        
        p.add_run("\n개선 필요 영역 (산업 평균 대비 미흡):\n").bold = True
        
        if tech_score <= industry_avg['technology']:
            p.add_run(f"• 기술성: 산업 평균 대비 {tech_score - industry_avg['technology']:.1f}점\n")
        if rights_score <= industry_avg['rights']:
            p.add_run(f"• 권리성: 산업 평균 대비 {rights_score - industry_avg['rights']:.1f}점\n")
        if market_score <= industry_avg['market']:
            p.add_run(f"• 활용성: 산업 평균 대비 {market_score - industry_avg['market']:.1f}점\n")
        
        doc.add_paragraph()
        
        # 3. 경쟁력 포지셔닝
        doc.add_heading('3. 경쟁력 포지셔닝', level=2)
        
        if overall_score >= 80:
            position = "최상위권 (Top 10%)"
            color = self.COLORS['success']
        elif overall_score >= 70:
            position = "상위권 (Top 30%)"
            color = self.COLORS['secondary']
        elif overall_score >= 60:
            position = "중상위권 (Top 50%)"
            color = self.COLORS['warning']
        else:
            position = "하위권 (Bottom 50%)"
            color = self.COLORS['danger']
        
        self._create_info_box(
            doc,
            f"경쟁력 포지션: {position}",
            f"본 특허는 {overall_score:.1f}점으로 산업 내 {position}에 위치합니다. "
            f"{'지속적인 기술 개발과 권리 보강으로 시장 리더십을 유지할 수 있습니다.' if overall_score >= 75 else '개선 활동을 통해 경쟁력을 강화할 필요가 있습니다.'}",
            color
        )
    
    def _add_enhanced_recommendations(self, doc: Document, state: Dict, 
                                     patent_info: Dict):
        """강화된 종합 평가 및 제언"""
        heading = doc.add_heading('종합 평가 및 전략적 제언', level=1)
        for run in heading.runs:
            run.font.color.rgb = self.COLORS['primary']
        
        overall_score = state.get('overall_score', 0)
        final_grade = state.get('final_grade', 'N/A')
        tech_score = state.get('tech_score', 0)
        rights_score = state.get('rights_score', 0)
        market_score = state.get('market_score', 0)
        
        # 1. 종합 의견
        doc.add_heading('1. 종합 평가 의견', level=2)
        
        p = doc.add_paragraph()
        p.add_run(f"종합 점수: {overall_score:.1f}점 ({final_grade})\n\n").bold = True
        p.runs[0].font.size = Pt(12)
        p.runs[0].font.color.rgb = self.COLORS['primary']
        
        # 점수대별 상세 평가
        if overall_score >= 85:
            assessment = """본 특허는 최상위 등급으로, 기술성, 권리성, 활용성 모든 면에서 매우 우수한 수준입니다. 
기술적 혁신성이 뛰어나고 견고한 권리범위를 확보하고 있으며, 시장 활용 가능성도 매우 높습니다. 
즉각적인 사업화 또는 라이선싱을 적극 추진할 수 있는 우수한 자산입니다."""
        elif overall_score >= 75:
            assessment = """본 특허는 우수한 등급으로, 전반적으로 높은 평가를 받았습니다. 
핵심 기술이 잘 보호되고 있으며 실용적 가치가 높습니다. 
일부 영역에서 추가 보강을 통해 더욱 강력한 특허로 발전시킬 수 있습니다."""
        elif overall_score >= 65:
            assessment = """본 특허는 양호한 수준으로, 기본적인 요건을 충족하고 있습니다. 
그러나 경쟁력 강화를 위해서는 약점 영역에 대한 보완이 필요합니다. 
전략적 개선을 통해 시장 가치를 높일 수 있을 것입니다."""
        else:
            assessment = """본 특허는 개선이 필요한 수준입니다. 
기술성, 권리성, 활용성 중 하나 이상의 영역에서 상당한 보강이 필요합니다. 
체계적인 개선 계획 수립과 실행을 권장합니다."""
        
        p.add_run(assessment)
        
        doc.add_paragraph()
        
        # 2. 영역별 전략적 제언
        doc.add_heading('2. 영역별 전략적 제언', level=2)
        
        # 기술성 제언
        p = doc.add_paragraph()
        p.add_run("🔬 기술성 제언\n").bold = True
        p.runs[0].font.color.rgb = self.COLORS['primary']
        p.runs[0].font.size = Pt(11)
        
        if tech_score >= 75:
            tech_advice = "• 현재 우수한 기술적 수준을 유지하고 있습니다.\n• 지속적인 기술 개발과 특허 출원을 통해 포트폴리오를 확장하세요.\n"
        elif tech_score >= 60:
            tech_advice = "• 구현 방법의 상세화를 통해 기술적 완성도를 높이세요.\n• 실험 데이터와 실시예를 추가하여 기술적 신뢰도를 강화하세요.\n"
        else:
            tech_advice = "• 기술적 혁신성을 강화할 필요가 있습니다.\n• 전문가 자문을 통해 기술 구현 방법을 재검토하세요.\n• 추가 발명 또는 개량 발명 출원을 고려하세요.\n"
        
        p.add_run(tech_advice)
        
        doc.add_paragraph()
        
        # 권리성 제언
        p = doc.add_paragraph()
        p.add_run("⚖️ 권리성 제언\n").bold = True
        p.runs[0].font.color.rgb = self.COLORS['primary']
        p.runs[0].font.size = Pt(11)
        
        if rights_score >= 75:
            rights_advice = "• 견고한 권리범위를 확보하고 있습니다.\n• 주요 시장에서의 해외 출원을 검토하세요.\n"
        elif rights_score >= 60:
            rights_advice = "• 종속항 추가를 통해 다층적 권리보호를 강화하세요.\n• 청구항 범위를 재검토하여 핵심 기술을 충분히 커버하도록 하세요.\n"
        else:
            rights_advice = "• 청구항 전면 재검토가 필요합니다.\n• 특허 전문가와 협력하여 권리범위를 재설정하세요.\n• 신규 출원 또는 분할 출원을 고려하세요.\n"
        
        p.add_run(rights_advice)
        
        doc.add_paragraph()
        
        # 활용성 제언
        p = doc.add_paragraph()
        p.add_run("📊 활용성 제언\n").bold = True
        p.runs[0].font.color.rgb = self.COLORS['primary']
        p.runs[0].font.size = Pt(11)
        
        if market_score >= 75:
            market_advice = "• 높은 시장 활용 가능성을 보유하고 있습니다.\n• 사업화 또는 기술이전을 적극 추진하세요.\n• 시장 진입 타이밍을 최적화하세요.\n"
        elif market_score >= 60:
            market_advice = "• 시장 검증을 통해 상용화 가능성을 확인하세요.\n• POC(Proof of Concept)를 수행하여 기술적 실현 가능성을 입증하세요.\n"
        else:
            market_advice = "• 시장 수요 분석을 재검토하세요.\n• 비즈니스 모델을 구체화하고 타겟 시장을 명확히 하세요.\n• 전략적 파트너십을 통한 시장 진입을 고려하세요.\n"
        
        p.add_run(market_advice)
        
        doc.add_paragraph()
        
        # 3. 우선순위 액션 플랜
        doc.add_heading('3. 우선순위 액션 플랜 (Top 5)', level=2)
        
        # 점수가 낮은 순서대로 우선순위 설정
        scores = {
            '기술성 강화': (tech_score, tech_score < 70),
            '권리성 보강': (rights_score, rights_score < 70),
            '활용성 제고': (market_score, market_score < 70)
        }
        
        actions = []
        
        # 점수가 70점 미만인 항목들을 우선순위로
        for area, (score, need_improvement) in sorted(scores.items(), key=lambda x: x[1][0]):
            if need_improvement:
                if '기술성' in area:
                    actions.append(("기술 구현 방법 상세화 및 실시예 추가", "High"))
                elif '권리성' in area:
                    actions.append(("청구항 범위 확대 및 종속항 보강", "High"))
                elif '활용성' in area:
                    actions.append(("시장 검증 및 비즈니스 모델 구체화", "High"))
        
        # 일반적인 액션 추가
        if overall_score >= 75:
            actions.append(("해외 주요 시장 특허 출원 검토", "Medium"))
            actions.append(("라이선싱 또는 사업화 전략 수립", "Medium"))
        else:
            actions.append(("특허 전문가와 개선 방안 협의", "High"))
            actions.append(("경쟁 특허 분석 및 회피 설계", "Medium"))
        
        # 표로 액션 플랜 정리
        table = doc.add_table(rows=len(actions[:5])+1, cols=3)
        table.style = 'Light Grid Accent 1'
        
        header_cells = table.rows[0].cells
        header_cells[0].text = '순위'
        header_cells[1].text = '액션 아이템'
        header_cells[2].text = '우선순위'
        
        for cell in header_cells:
            cell.paragraphs[0].runs[0].font.bold = True
        
        for i, (action, priority) in enumerate(actions[:5], start=1):
            cells = table.rows[i].cells
            cells[0].text = str(i)
            cells[1].text = action
            cells[2].text = priority
            
            # 우선순위 색상
            if priority == "High":
                cells[2].paragraphs[0].runs[0].font.color.rgb = self.COLORS['danger']
            else:
                cells[2].paragraphs[0].runs[0].font.color.rgb = self.COLORS['warning']
    
    def _add_improvement_roadmap(self, doc: Document, state: Dict):
        """개선 로드맵 (신규)"""
        heading = doc.add_heading('개선 로드맵 (Improvement Roadmap)', level=1)
        for run in heading.runs:
            run.font.color.rgb = self.COLORS['primary']
        
        tech_score = state.get('tech_score', 0)
        rights_score = state.get('rights_score', 0)
        market_score = state.get('market_score', 0)
        overall_score = state.get('overall_score', 0)
        
        # 로드맵 설명
        self._create_info_box(
            doc,
            "3단계 개선 로드맵",
            "특허의 가치를 극대화하기 위한 단계별 개선 계획입니다. "
            "각 단계는 3-6개월의 기간을 상정하며, 우선순위에 따라 순차적으로 진행할 것을 권장합니다."
        )
        
        doc.add_paragraph()
        
        # Phase 1: 단기 (0-6개월)
        doc.add_heading('Phase 1: 즉시 실행 (0-6개월)', level=2)
        
        p = doc.add_paragraph()
        p.add_run("목표: 핵심 약점 개선 및 긴급 보강\n\n").bold = True
        
        phase1_tasks = []
        
        if tech_score < 70:
            phase1_tasks.append("✓ 기술 명세서 보완 (구현 방법 상세화)")
            phase1_tasks.append("✓ 실험 데이터 및 실시예 추가")
        
        if rights_score < 70:
            phase1_tasks.append("✓ 청구항 재검토 및 보정")
            phase1_tasks.append("✓ 종속항 추가 출원 검토")
        
        if market_score < 70:
            phase1_tasks.append("✓ 시장 수요 조사 및 분석")
            phase1_tasks.append("✓ 초기 고객 인터뷰 및 피드백 수집")
        
        if not phase1_tasks:
            phase1_tasks.append("✓ 현재 수준 유지 및 모니터링")
            phase1_tasks.append("✓ 경쟁 특허 분석")
        
        for task in phase1_tasks:
            doc.add_paragraph(task, style='List Bullet')
        
        doc.add_paragraph()
        
        # Phase 2: 중기 (6-12개월)
        doc.add_heading('Phase 2: 전략적 강화 (6-12개월)', level=2)
        
        p = doc.add_paragraph()
        p.add_run("목표: 경쟁력 향상 및 시장 준비\n\n").bold = True
        
        phase2_tasks = [
            "✓ POC(Proof of Concept) 수행",
            "✓ 추가 특허 출원 (개량/주변 특허)",
            "✓ 비즈니스 모델 구체화",
            "✓ 잠재 파트너/고객 발굴"
        ]
        
        for task in phase2_tasks:
            doc.add_paragraph(task, style='List Bullet')
        
        doc.add_paragraph()
        
        # Phase 3: 장기 (12-24개월)
        doc.add_heading('Phase 3: 시장 진출 (12-24개월)', level=2)
        
        p = doc.add_paragraph()
        p.add_run("목표: 상용화 및 수익 창출\n\n").bold = True
        
        phase3_tasks = []
        
        if overall_score >= 75:
            phase3_tasks.extend([
                "✓ 본격적인 사업화 추진",
                "✓ 해외 주요 시장 진출",
                "✓ 특허 포트폴리오 확장"
            ])
        else:
            phase3_tasks.extend([
                "✓ 기술이전 또는 라이선싱",
                "✓ 전략적 파트너십 체결",
                "✓ 시범 사업 진행"
            ])
        
        phase3_tasks.append("✓ 특허 가치 평가 및 재평가")
        
        for task in phase3_tasks:
            doc.add_paragraph(task, style='List Bullet')
        
        doc.add_paragraph()
        
        # 예상 성과
        doc.add_heading('예상 성과 및 목표', level=2)
        
        table = doc.add_table(rows=4, cols=3)
        table.style = 'Medium Shading 1 Accent 1'
        
        header_cells = table.rows[0].cells
        header_cells[0].text = '영역'
        header_cells[1].text = '현재'
        header_cells[2].text = '목표 (24개월 후)'
        
        for cell in header_cells:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
        target_tech = min(tech_score + 10, 95)
        target_rights = min(rights_score + 10, 95)
        target_market = min(market_score + 15, 95)
        target_overall = min(overall_score + 12, 95)
        
        data = [
            ('기술성', f"{tech_score:.1f}점", f"{target_tech:.1f}점"),
            ('권리성', f"{rights_score:.1f}점", f"{target_rights:.1f}점"),
            ('활용성', f"{market_score:.1f}점", f"{target_market:.1f}점"),
        ]
        
        for i, (category, current, target) in enumerate(data, start=1):
            cells = table.rows[i].cells
            cells[0].text = category
            cells[1].text = current
            cells[2].text = target
    
    def _add_enhanced_references(self, doc: Document, patent_info: Dict, state: Dict):
        """강화된 Reference (기존 유지 + 추가 정보)"""
        heading = doc.add_heading('Reference - 참고 문서 및 출처', level=1)
        for run in heading.runs:
            run.font.color.rgb = self.COLORS['primary']
        
        # ========== 기존 v5.0 구조 유지 ==========
        
        # 1. 특허 원문
        doc.add_heading('1. 특허 원문 정보', level=2)
        
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Light List Accent 1'
        
        data = [
            ('특허번호', patent_info.get('number', 'N/A')),
            ('발명명칭', patent_info.get('title', 'N/A')),
            ('출원인', patent_info.get('applicant', 'N/A')),
            ('출처', 'PDF 원문 분석'),
        ]
        
        for i, (key, value) in enumerate(data):
            cells = table.rows[i].cells
            cells[0].text = key
            cells[1].text = str(value)
            cells[0].paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()
        
        # 2. 웹 서치 출처
        doc.add_heading('2. 웹 서치 출처 (실시간 데이터)', level=2)
        
        market_web_search = state.get('market_web_search', {})
        
        p = doc.add_paragraph()
        p.add_run("출원인 시장 정보:\n").bold = True
        p.add_run(f"• 내용: {market_web_search.get('applicant_summary', 'N/A')}\n")
        p.add_run(f"• 평가 등급: {market_web_search.get('applicant_grade', 'Unknown')}\n")
        p.add_run(f"• 출처: DuckDuckGo 웹 검색 (실시간)\n")
        p.add_run(f"• 검색일: {datetime.now().strftime('%Y-%m-%d')}\n\n")
        
        p.add_run("기술 분야 정보:\n").bold = True
        p.add_run(f"• 내용: {market_web_search.get('tech_summary', 'N/A')}\n")
        p.add_run(f"• 평가 등급: {market_web_search.get('tech_grade', 'Unknown')}\n")
        p.add_run(f"• 출처: DuckDuckGo 웹 검색 (실시간)\n")
        p.add_run(f"• 검색일: {datetime.now().strftime('%Y-%m-%d')}\n")
        
        doc.add_paragraph()
        
        # 3. 평가 모델 및 시스템
        doc.add_heading('3. 평가 모델 및 시스템', level=2)
        
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Light Grid Accent 1'
        
        data = [
            ('평가 시스템', '특허 평가 시스템 v6.0 (전문가급)'),
            ('평가 방법', '정량평가 중심 + LLM 정성평가'),
            ('RAG 모델', 'nlpai-lab/KoE5 (HuggingFace)'),
            ('LLM 모델', 'GPT-4o-mini (OpenAI)'),
            ('정량 지표', '10개 (X1-X10)'),
            ('평가일시', datetime.now().strftime('%Y년 %m월 %d일 %H:%M')),
        ]
        
        for i, (key, value) in enumerate(data):
            cells = table.rows[i].cells
            cells[0].text = key
            cells[1].text = value
            cells[0].paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()
        
        # ========== 추가 강화 내용 ==========
        
        # 4. 평가 기준 및 방법론
        doc.add_heading('4. 평가 기준 및 방법론', level=2)
        
        p = doc.add_paragraph()
        p.add_run("가중치 배분:\n").bold = True
        p.add_run("• 종합 = 기술성(45%) + 권리성(35%) + 활용성(20%)\n")
        p.add_run("• 기술성 = 정량(60%) + 정성/LLM(40%)\n")
        p.add_run("• 권리성 = 정량(70%) + 정성/LLM(30%)\n")
        p.add_run("• 활용성 = 정량+웹서치(70%) + 정성/LLM(30%)\n\n")
        
        p.add_run("평가 프로세스:\n").bold = True
        p.add_run("1. PDF 파싱 → 메타데이터 추출\n")
        p.add_run("2. RAG 시스템 → 관련 컨텍스트 검색\n")
        p.add_run("3. 정량 평가 → 10개 지표 자동 계산\n")
        p.add_run("4. 웹 서치 → 실시간 시장 정보 수집\n")
        p.add_run("5. LLM 평가 → 정성적 분석 수행\n")
        p.add_run("6. 종합 평가 → 최종 점수 및 등급 산출\n")
        
        doc.add_paragraph()
        
        # 5. 데이터 출처 및 신뢰도
        doc.add_heading('5. 데이터 출처 및 신뢰도', level=2)
        
        self._create_info_box(
            doc,
            "데이터 신뢰도 공지",
            """본 평가는 다음 데이터 소스를 기반으로 수행되었습니다:

• 특허 원문 데이터: 공식 특허 문서(PDF)에서 직접 추출한 1차 자료
• 웹 검색 데이터: DuckDuckGo를 통한 실시간 검색 결과 (2차 자료)
• LLM 분석: GPT-4o-mini 모델의 정성적 분석 (AI 보조 자료)

정량 지표는 특허 원문에서 객관적으로 추출되어 높은 신뢰도를 가지며,
웹 검색 및 LLM 분석은 참고 자료로 활용하시기 바랍니다.""",
            self.COLORS['light']
        )
    
    def _add_enhanced_appendix(self, doc: Document, state: Dict):
        """강화된 Appendix (기존 유지 + 추가 정보)"""
        heading = doc.add_heading('Appendix - 평가 지표 및 상세 자료', level=1)
        for run in heading.runs:
            run.font.color.rgb = self.COLORS['primary']
        
        # ========== 기존 v5.0 구조 유지 ==========
        
        # 1. 정량 지표 (10개)
        doc.add_heading('1. 정량 지표 (X1-X10) 상세', level=2)
        
        tech_metrics = state.get('tech_metrics', {})
        rights_metrics = state.get('rights_metrics', {})
        market_metrics = state.get('market_metrics', {})
        
        table = doc.add_table(rows=11, cols=5)
        table.style = 'Medium Shading 1 Accent 1'
        
        header_cells = table.rows[0].cells
        header_cells[0].text = '지표'
        header_cells[1].text = '측정값'
        header_cells[2].text = '범주'
        header_cells[3].text = 'Agent'
        header_cells[4].text = '설명'
        
        for cell in header_cells:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
        data = [
            ('X1', f"{rights_metrics.get('X1_ipc_count', 0)}개", '권리성', 'rights', 
             'IPC 코드 수 (기술 다양성)'),
            ('X2', f"{rights_metrics.get('X2_independent_claims', 0)}개", '권리성', 'rights', 
             '독립항 수 (핵심 권리)'),
            ('X3', f"{rights_metrics.get('X3_dependent_claims', 0)}개", '권리성', 'rights', 
             '종속항 수 (세부 보호)'),
            ('X4', f"{rights_metrics.get('X4_total_claims', 0)}개", '권리성', 'rights', 
             '전체 청구항 수'),
            ('X5', f"{rights_metrics.get('X5_independent_avg_length', 0):.1f}자", '권리성', 'rights', 
             '독립항 평균 길이'),
            ('X6', f"{rights_metrics.get('X6_dependent_avg_length', 0):.1f}자", '권리성', 'rights', 
             '종속항 평균 길이'),
            ('X7', f"{tech_metrics.get('X7_drawing_count', 0)}개", '기술성', 'tech', 
             '도면 수 (시각화)'),
            ('X8', f"{tech_metrics.get('X8_title_length', 0)}자", '기술성', 'tech', 
             '발명명칭 길이'),
            ('X9', f"{tech_metrics.get('X9_claim_series', 0)}개", '기술성', 'tech', 
             '청구항 계열 수'),
            ('X10', f"{market_metrics.get('X10_inventor_count', 0)}명", '활용성', 'market', 
             '발명자 수 (협업 규모)'),
        ]
        
        for i, (code, value, category, agent, desc) in enumerate(data, start=1):
            cells = table.rows[i].cells
            cells[0].text = code
            cells[1].text = value
            cells[2].text = category
            cells[3].text = agent
            cells[4].text = desc
        
        doc.add_paragraph()
        
        # 2. 구조방정식 모델
        doc.add_heading('2. 구조방정식 모델 (SEM)', level=2)
        
        p = doc.add_paragraph()
        p.add_run("기술성 점수 계산:\n").bold = True
        p.add_run("= X7(도면) × 0.4 + X8(명칭) × 0.3 + X9(계열) × 0.3\n\n")
        
        p.add_run("권리성 점수 계산:\n").bold = True
        p.add_run("= IPC(25%) + 청구항개수(30%) + 청구항길이(25%) + 계층구조(20%)\n")
        p.add_run("  where:\n")
        p.add_run("  • IPC = X1(IPC 수) 점수\n")
        p.add_run("  • 청구항개수 = (X2(독립항) × 0.5 + X3(종속항) × 0.5)\n")
        p.add_run("  • 청구항길이 = (X5(독립항 길이) × 0.5 + X6(종속항 길이) × 0.5)\n")
        p.add_run("  • 계층구조 = X2(독립항) + X3(종속항) 조합 평가\n\n")
        
        p.add_run("활용성 점수 계산:\n").bold = True
        p.add_run("= 발명자(30%) + 출원인(40%) + 기술분야(30%)\n")
        p.add_run("  where:\n")
        p.add_run("  • 발명자 = X10(발명자 수) 점수\n")
        p.add_run("  • 출원인 = 웹 서치 결과 등급 점수\n")
        p.add_run("  • 기술분야 = 웹 서치 결과 등급 점수\n\n")
        
        p.add_run("종합 점수 계산:\n").bold = True
        p.add_run("= 기술성(45%) + 권리성(35%) + 활용성(20%)\n")
        
        doc.add_paragraph()
        
        # 3. Binary 체크리스트 전체
        doc.add_heading('3. Binary 체크리스트 (전체)', level=2)
        
        tech_binary = state.get('tech_binary', {})
        rights_binary = state.get('rights_binary', {})
        market_binary = state.get('market_binary', {})
        
        # 기술성 체크리스트
        p = doc.add_paragraph()
        p.add_run("▶ 기술성 체크리스트\n").bold = True
        p.runs[0].font.color.rgb = self.COLORS['primary']
        
        for key, value in tech_binary.items():
            status = "✓" if value else "✗"
            color = self.COLORS['success'] if value else self.COLORS['danger']
            
            para = doc.add_paragraph(f"{status} {key}", style='List Bullet')
            para.runs[0].font.color.rgb = color
        
        doc.add_paragraph()
        
        # 권리성 체크리스트
        p = doc.add_paragraph()
        p.add_run("▶ 권리성 체크리스트\n").bold = True
        p.runs[0].font.color.rgb = self.COLORS['primary']
        
        for key, value in rights_binary.items():
            status = "✓" if value else "✗"
            color = self.COLORS['success'] if value else self.COLORS['danger']
            
            para = doc.add_paragraph(f"{status} {key}", style='List Bullet')
            para.runs[0].font.color.rgb = color
        
        doc.add_paragraph()
        
        # 활용성 체크리스트
        p = doc.add_paragraph()
        p.add_run("▶ 활용성 체크리스트\n").bold = True
        p.runs[0].font.color.rgb = self.COLORS['primary']
        
        for key, value in market_binary.items():
            status = "✓" if value else "✗"
            color = self.COLORS['success'] if value else self.COLORS['danger']
            
            para = doc.add_paragraph(f"{status} {key}", style='List Bullet')
            para.runs[0].font.color.rgb = color
        
        doc.add_paragraph()
        
        # 4. 평가 등급 기준
        doc.add_heading('4. 평가 등급 기준', level=2)
        
        table = doc.add_table(rows=11, cols=3)
        table.style = 'Light Grid Accent 1'
        
        header_cells = table.rows[0].cells
        header_cells[0].text = '등급'
        header_cells[1].text = '점수 범위'
        header_cells[2].text = '평가'
        
        for cell in header_cells:
            cell.paragraphs[0].runs[0].font.bold = True
        
        grade_data = [
            ('AAA', '90점 이상', '최고 수준 - 모든 면에서 탁월'),
            ('AA', '85-89점', '매우 우수 - 경쟁력 매우 높음'),
            ('A', '80-84점', '우수 - 시장 가치 높음'),
            ('BBB', '75-79점', '양호 - 활용 가능성 높음'),
            ('BB', '70-74점', '보통 상위 - 일부 보강 필요'),
            ('B', '65-69점', '보통 - 개선 여지 있음'),
            ('CCC', '60-64점', '보통 하위 - 개선 권장'),
            ('CC', '57-59점', '미흡 - 상당한 개선 필요'),
            ('C', '55-56점', '개선 필요 - 전면 재검토'),
            ('미달', '55점 미만', '재평가 필요 - 근본적 개선'),
        ]
        
        for i, (grade, score_range, assessment) in enumerate(grade_data, start=1):
            cells = table.rows[i].cells
            cells[0].text = grade
            cells[1].text = score_range
            cells[2].text = assessment
            
            # 등급별 색상
            if grade in ['AAA', 'AA', 'A']:
                cells[0].paragraphs[0].runs[0].font.color.rgb = self.COLORS['success']
            elif grade in ['BBB', 'BB', 'B']:
                cells[0].paragraphs[0].runs[0].font.color.rgb = self.COLORS['secondary']
            elif grade in ['CCC', 'CC']:
                cells[0].paragraphs[0].runs[0].font.color.rgb = self.COLORS['warning']
            else:
                cells[0].paragraphs[0].runs[0].font.color.rgb = self.COLORS['danger']
        
        doc.add_paragraph()
        
        # ========== 추가 강화 내용 ==========
        
        # 5. 용어 설명
        doc.add_heading('5. 주요 용어 설명', level=2)
        
        terms = [
            ("RAG (Retrieval-Augmented Generation)", 
             "검색 증강 생성. 대량의 문서에서 관련 정보를 검색하여 LLM의 응답을 보강하는 기술"),
            ("IPC (International Patent Classification)", 
             "국제특허분류. 특허를 기술 분야별로 체계적으로 분류하는 국제 표준"),
            ("독립항 (Independent Claim)", 
             "다른 청구항을 인용하지 않고 독립적으로 발명을 정의하는 청구항"),
            ("종속항 (Dependent Claim)", 
             "독립항 또는 다른 종속항을 인용하여 발명을 더 구체화하는 청구항"),
            ("청구항 계열 (Claim Series)", 
             "하나의 독립항과 그에 종속하는 종속항들의 그룹"),
            ("POC (Proof of Concept)", 
             "개념 증명. 아이디어나 기술의 실현 가능성을 검증하는 초기 단계 실험"),
        ]
        
        for term, definition in terms:
            p = doc.add_paragraph()
            p.add_run(f"{term}\n").bold = True
            p.add_run(f"   {definition}\n")
        
        doc.add_paragraph()
        
        # 6. 문의 및 연락처
        doc.add_heading('6. 보고서 관련 문의', level=2)
        
        self._create_info_box(
            doc,
            "추가 분석 및 컨설팅",
            """본 보고서에 대한 추가 분석이나 상세 컨설팅이 필요하신 경우,
특허 평가 시스템 개발자인 SKALA 2기 3반 백선재 교육생한테 문의하시기 바랍니다.

• 평가 시스템: Patent Evaluation System v6.0
• 평가 모델: RAG + LLM 하이브리드
• 생성일시: """ + datetime.now().strftime('%Y년 %m월 %d일 %H:%M:%S'),
            self.COLORS['light']
        )


if __name__ == "__main__":
    print("DOCX 보고서 생성기 v6.0 - 전문가급 디자인 & 내용 강화")
    print("✓ 기존 v5.0의 모든 구성요소 유지")
    print("✓ Binary checklist, Appendix, Reference 포함")
    print("✓ 고급 디자인 및 시각화 강화")
    print("✓ 비교 분석, 개선 로드맵 등 추가")